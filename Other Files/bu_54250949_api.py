
import os
import faiss
import pickle
import json
import base64
import datetime
import re
import numpy as np
import requests
import textwrap
import openai
import zipfile
from io import BytesIO
from flask_cors import CORS
from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from zoneinfo import ZoneInfo

# ✅ Unzip chunks.zip once at startup
import os
import zipfile

# Path to zip file and extraction folder (relative to api.py)
base_dir = os.path.dirname(os.path.abspath(__file__))
zip_path = os.path.join(base_dir, "chunks.zip")
chunks_dir = os.path.join(base_dir, "data/accounting")

# Unzip only if not already unzipped
sample_chunk = os.path.join(chunks_dir, "Check when you must use the VAT domestic reverse charge for building and construction services - GOV.UK_chunk_2.txt")

if os.path.exists(zip_path) and not os.path.exists(sample_chunk):
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(chunks_dir)
            print("✅ Unzipped chunks.zip to data/accounting/")
    except Exception as e:
        print(f"❌ Failed to unzip chunks.zip: {e}")

__version__ = "v1.0.7-test"
print(f"🚀 API Version: {__version__}")


# ✅ Unzip chunks.zip once at startup
zip_path = "data/accounting/chunks.zip"
sample_txt = "data/accounting/Check when you must use the VAT domestic reverse charge for building and construction services - GOV.UK_chunk_2.txt"

if os.path.exists(zip_path) and not os.path.exists(sample_txt):
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall("data/accounting")
            print("✅ Unzipped chunks.zip to data/accounting/")
    except Exception as e:
        print(f"❌ Failed to unzip chunks.zip: {e}")

__version__ = "v1.0.7-test"
print(f"🚀 API Version: {__version__}")

from openai import OpenAI
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = Flask(__name__)
CORS(app, origins=["https://www.aivs.uk"])

@app.route("/", methods=["GET"])
def home():
    return "✅ Business API is running", 200

@app.after_request
def apply_cors_headers(response):
    response.headers.add("Access-Control-Allow-Origin", "https://www.aivs.uk")
    response.headers.add("Access-Control-Allow-Headers", "Content-Type")
    response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
    return response

@app.route("/ping", methods=["POST", "OPTIONS"])
def ping():
    if request.method == "OPTIONS":
        return '', 204
    return jsonify({"message": "pong"})

# Load FAISS index
try:
    faiss_index = faiss.read_index("data/accounting/accounting.index")
    #with open("data/accounting/metadata.pkl", "rb") as f:
        #metadata = pickle.load(f)
    with open("data/accounting/accounting_metadata.json", "r", encoding="utf-8") as f:
        metadata = json.load(f)
    print("✅ Accounting FAISS index and metadata loaded.")
except Exception as e:
    faiss_index = None
    metadata = []
    print("⚠️ Failed to load Accounting FAISS index:", str(e))

# PART 2

def ask_gpt_with_context(data, context):
    query = data.get("query", "")
    job_title = data.get("job_title", "Not specified")
    seniority_level = data.get("seniority_level", "Not specified") 
    timeline = data.get("timeline", "Not specified")
    discipline = data.get("discipline", "Not specified")
    site = data.get("site", "Not specified")
    funnel_1 = data.get("funnel_1", "Not specified")
    funnel_2 = data.get("funnel_2", "Not specified")
    funnel_3 = data.get("funnel_3", "Not specified")

    prompt = f"""
You are responding to a professional business query via a secure reporting system.

All responses must:
- Be based on correct UK financial standards, accounting regulations, business risk practices, or strategic management theory.
- Use British English spelling and tone.

### Enquiry:
\"{query}\"

### Context from FAISS Index:
{context}

### Additional Focus:
- Support Need: {funnel_1}
- Current Status: {funnel_2}
- Follow-Up Expectation: {funnel_3}

### Your Task:
Please generate a structured response that includes:
1. **Client Reply**
2. **Action Sheet**
3. **Policy or Standard Notes**
"""
    return generate_reviewed_response(prompt, discipline)

def generate_reviewed_response(prompt, discipline):
    print("📢 Sending initial GPT prompt...")
    completion = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=1800
    )
    initial_response = completion.choices[0].message.content.strip()
    print(f"📏 Initial GPT response length: {len(initial_response)} characters")

    if len(initial_response) > 1500:
        print("⚡ Skipping review — using initial GPT response directly.")
        return initial_response

    print("🔄 Reviewing GPT response...")
    initial_response = re.sub(r'(Best regards,|Yours sincerely,|Kind regards,)[\s\S]*$', '', initial_response, flags=re.IGNORECASE).strip()
    stripped_response = initial_response.split("### Context from FAISS Index:")[0].strip()
    stripped_response = stripped_response[:2000]

    review_prompt = textwrap.dedent(f"""\
    Please clean and improve the following structured response while maintaining professional tone and factual accuracy.
    --- START RESPONSE ---
    {stripped_response}
    --- END RESPONSE ---
    """)

    review_completion = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": review_prompt}],
        temperature=0,
        max_tokens=700,
        timeout=15
    )

    reviewed_response = review_completion.choices[0].message.content.strip()
    print(f"✅ Reviewed response length: {len(reviewed_response)} characters")
    return reviewed_response

# PART 3

def send_email_mailjet(to_emails, subject, body_text, doc_buffer, full_name=None, supervisor_name=None):
    MAILJET_API_KEY = os.getenv("MJ_APIKEY_PUBLIC")
    MAILJET_SECRET_KEY = os.getenv("MJ_APIKEY_PRIVATE")

    messages = []

    for recipient in to_emails:
        role = recipient["Name"]
        email = recipient["Email"]

        text_body = f"This document was generated following a query submitted by {full_name}. Please file or follow up according to internal procedures."

        messages.append({
            "From": {
                "Email": "noreply@securemaildrop.uk",
                "Name": "Secure Maildrop"
            },
            "To": [{"Email": email, "Name": role}],
            "Subject": subject,
            "TextPart": text_body,
            "HTMLPart": f"<pre>{text_body}</pre>",
            "Attachments": [
                {
                    "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "Filename": f"{full_name.replace(' ', '_')}_Response.docx",
                    "Base64Content": base64.b64encode(doc_buffer.read()).decode()
                }
            ]
        })

    response = requests.post(
        "https://api.mailjet.com/v3.1/send",
        auth=(MAILJET_API_KEY, MAILJET_SECRET_KEY),
        json={"Messages": messages}
    )
    print(f"📤 Mailjet status: {response.status_code}")
    print(response.json())
    return response.status_code, response.json()

@app.route("/generate", methods=["POST"])
def generate_response():
    print("📥 /generate route hit")

    
    try:
        data = request.get_json()
        print("🔎 Payload received:", data)
    except Exception as e:
        print("❌ Error parsing JSON:", e)
        return jsonify({"error": "Invalid JSON input"}), 400

    query_text = data.get("query")
    full_name = data.get("full_name", "User")
    user_email = data.get("user_email")
    supervisor_email = data.get("supervisor_email")
    hr_email = data.get("hr_email")
    supervisor_name = data.get("supervisor_name", "Supervisor")
    timestamp = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")

    if faiss_index:
        query_vector = client.embeddings.create(
            input=[query_text.replace("\n", " ")],
            model="text-embedding-3-small"
        ).data[0].embedding

        D, I = faiss_index.search(np.array([query_vector]).astype("float32"), 2)
        # ✅ Log matched chunk filenames
        print("📂 Matched chunk files:")
        for i in I[0]:
           key = str(i)
           chunk_file = metadata.get(key, {}).get("chunk_file")
           if chunk_file:
              print(f" - {chunk_file}")

        matched_chunks = []
        for i in I[0]:
            key = str(i)
            if "chunk_file" in metadata.get(key, {}):
                chunk_file = metadata[key]["chunk_file"]
                with open(f"data/accounting/{chunk_file}", "r", encoding="utf-8") as f:
                    matched_chunks.append(f.read().strip())
            else:
                print(f"⚠️ Missing 'chunk_file' for index {i} in metadata.")
         
        context = "\n\n---\n\n".join(matched_chunks)

    else:
        context = "Policy lookup not available (FAISS index not loaded)."

    answer = ask_gpt_with_context(data, context)
    answer = re.sub(r"### ORIGINAL QUERY\s*[\r\n]+.*?(?=###|\Z)", "", answer, flags=re.IGNORECASE | re.DOTALL).strip()
    print(f"🧠 GPT answer: {answer[:80]}...")

    discipline = data.get("discipline", "Not specified")
    discipline_folder = discipline.lower().replace(" ", "_")
    output_path = f"output/{discipline_folder}"
    os.makedirs(output_path, exist_ok=True)

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"RESPONSE FOR {full_name.upper()}")
    title_run.bold = True
    title_run.font.size = Pt(14)

    uk_time = datetime.datetime.now(ZoneInfo("Europe/London"))
    generated_datetime = uk_time.strftime("%d %B %Y at %H:%M:%S (%Z)")
    doc.add_paragraph(f"Generated: {generated_datetime}")

    # Add AI Response content (simple insert for now)
    doc.add_paragraph(answer)

    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    recipients = []
    if user_email:
        recipients.append({"Email": user_email, "Name": full_name})
    if supervisor_email:
        recipients.append({"Email": supervisor_email, "Name": supervisor_name})
    if hr_email:
        recipients.append({"Email": hr_email, "Name": "HR Department"})

    if not recipients:
        return jsonify({"error": "No valid email addresses provided."}), 400

    subject = f"AI Analysis for {full_name} - {timestamp}"
    body_text = f"This document was generated following a query submitted by {full_name}. Please file or follow up according to internal procedures."

    status, response = send_email_mailjet(
        to_emails=recipients,
        subject=subject,
        body_text=body_text,
        doc_buffer=doc_buffer,
        full_name=full_name,
        supervisor_name=supervisor_name
    )

    return jsonify({
        "status": "ok",
        "message": "✅ OpenAI-powered response generated, AI reviewed and email successfully sent.",
        "disclaimer": "This document was generated by AIVS Software Limited using AI assistance (OpenAI). Please review for accuracy and relevance before taking any formal action.",
        "context_preview": context[:200],
        "mailjet_status": status,
        "mailjet_response": response
    })

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)